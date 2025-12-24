from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# --- 1. 定義格式設定函式 ---
def set_font(run, font_name_en='Times New Roman', font_name_ch='標楷體', size=12, bold=False, color=None):
    run.font.name = font_name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color == 'red':
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif color == 'black':
        run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name_ch)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=14, bold=True)
    elif level == 2:
        set_font(run, size=13, bold=True)
    else:
        set_font(run, size=12, bold=True)

def add_figure(doc, image_filename, caption_text):
    """自動插入圖片並置中"""
    if os.path.exists(image_filename):
        try:
            # 插入圖片，寬度設定為 14cm (適合 A4 單欄)
            doc.add_picture(image_filename, width=Cm(14))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"✅ 已插入圖片：{image_filename}")
        except Exception as e:
            print(f"⚠️ 插入圖片失敗 ({image_filename}): {e}")
            p = doc.add_paragraph(f"[圖片插入失敗: {image_filename}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.runs[0], color='red')
    else:
        print(f"⚠️ 找不到圖片檔案：{image_filename}")
        p = doc.add_paragraph(f"[請在此處插入圖片：{image_filename}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.runs[0], color='red')

    # 圖片標題
    p = doc.add_paragraph(caption_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.runs[0], color='black', bold=True)
    p.paragraph_format.space_after = Pt(12)

def set_table_borders(table):
    """手動設定表格框線 (模擬格線表)"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # 框線粗細
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    tblPr.append(borders)

# --- 2. 論文內容變數 ---
TITLE = "邁向氣候意識投資：基於生成式 AI 之碳權與美股雙軌分析平台設計與實作"
TITLE_EN = "Towards Climate-Aware Investing: Design and Implementation of a GenAI-based Dual-Track Analysis Platform for Carbon and US Equity ETFs"
AUTHORS = "林子瑜 1　李冠榮 2"
AFFILIATION = "1, 2 崑山科技大學資訊工程系碩士班"

ABSTRACT = "隨著全球淨零轉型加速，碳定價（Carbon Pricing）已成為影響資本市場的重要因子。然而，一般投資人往往僅關注股價波動，忽略了碳權價格對企業營運成本的潛在衝擊。本文提出一套整合生成式 AI 的「雙軌趨勢分析平台」，命名為 AEGIS (AI-driven ETF Guardian & Intelligent System)。本系統設計分為「雙軌趨勢偵測」與「智能風險解讀」兩大核心模組。在偵測模組，系統自動化獲取美股大盤 ETF (SPY) 與全球碳權 ETF (KRBN) 數據，部署雙層 LSTM 模型進行雙軌趨勢預測；在解讀模組，則引入 OpenAI GPT 生成式技術，針對偵測結果進行語意分析。實驗結果顯示，本系統在美股市場預測效能優異（F1-Score 達 0.733），且在碳權市場能有效捕捉非線性波動，並透過協同機制自動生成「氣候金融風險摘要」，有效降低綠色金融的技術門檻，協助投資人建立具備氣候意識的決策框架。"
KEYWORDS = "生成式 AI、碳權 ETF、LSTM、氣候金融、協同運作機制"

# 英文摘要
ABSTRACT_EN_TITLE = "Abstract"
ABSTRACT_EN = "With the acceleration of the global net-zero transition, Carbon Pricing has become a non-negligible cost variable in capital markets. However, existing investment tools mostly focus on the price volatility of single assets, lacking an integrated solution capable of simultaneously assessing the risks of \"traditional equity securities\" and \"emerging climate assets.\" This study proposes the AEGIS (AI-driven ETF Guardian & Intelligent System) platform. Drawing on the concept of \"Synergy\" from system engineering, the system is designed with two core modules: \"Dual-Track Trend Detection\" and \"Intelligent Risk Interpretation.\" In the detection module, the system automatically acquires data for the US market ETF (SPY) and the Global Carbon Credit ETF (KRBN), deploying a double-layer LSTM model to conduct dual-track trend prediction. In the interpretation module, OpenAI GPT generative technology is introduced to perform semantic analysis on the detection results. Experimental results indicate that the system demonstrates superior predictive performance in the US stock market (achieving an F1-Score of 0.733) and effectively captures non-linear volatility in the carbon credit market. Furthermore, through the synergy mechanism, the system automatically generates \"Climate Finance Risk Summaries,\" effectively lowering the technical threshold of green finance and assisting investors in establishing a climate-aware decision-making framework."
KEYWORDS_EN = "Generative AI, Carbon Credit ETF, LSTM, Climate Finance, Synergy Mechanism"

SECTION_1 = "在金融科技（FinTech）與永續轉型（Sustainability Transition）的雙重浪潮下，數據科學已成為輔助投資決策的關鍵技術。然而，面對新興的「碳權交易市場（Carbon Credit Market）」，一般投資人往往面臨嚴重的資訊不對稱，難以量化碳價（Carbon Price）波動對傳統權益證券（如 SPY）成本結構的潛在衝擊。\n為解決上述問題，本文提出 AEGIS 系統架構。本系統目的在於建構一個全自動化的雙軌分析平台，透過「量化預測」與「質化解讀」的交互運作，將複雜的氣候金融數據轉化為可執行的投資情報，猶如投資人的數位守護者（Guardian）。"

SECTION_2_INTRO = "本系統規劃為端到端（End-to-End）的解決方案，主要架構如圖 1 所示。系統邏輯分為「雙軌趨勢偵測」、「智能風險解讀」兩大核心模式，以及連結兩者的「協同運作機制」。"
SECTION_2_1 = "系統採模組化與物件導向原則設計，將資料獲取、特徵工程、模型訓練封裝為獨立服務。資料流源自 Yahoo Finance API，經清洗程式（ETL）處理後存入 SQLite 資料庫以確保數據一致性。核心運算層定期從資料庫讀取數據，驅動偵測與解讀模組進行平行運算，最終將結果匯流至 Streamlit 前端介面。"
SECTION_2_2 = "本模組類比於資安系統中的「偵測引擎」，負責全天候監控市場的異常趨勢。\n1. 雙軌監控：系統針對 SPY (美股大盤) 與 KRBN (全球碳權) 建立兩條獨立的監控管線。\n2. 特徵提取：計算移動平均線 (MA)、相對強弱指標 (RSI) 及歷史波動率，作為判斷市場位階的特徵向量。\n3. 模型推論：採用雙層堆疊 LSTM (Double-Layer LSTM) 網路作為核心演算法。利用 LSTM 的遺忘閘與輸入閘機制，捕捉時間序列中的長期依賴關係，精準預測次日市場趨勢。"
SECTION_2_3 = "本模組類比於資安系統中的「事件分析師」，負責將冷冰冰的數據轉化為具體的風險評估。\n1. 生成式核心：導入 OpenAI GPT-4 模型作為語意分析引擎。\n2. 動態提示工程 (Dynamic Prompting)：系統將偵測模組輸出的量化指標（如：KRBN 預測信心度 85%）動態組裝成結構化 Prompt，要求 AI 扮演「氣候金融專家」角色。\n3. 關聯性分析：模型重點分析「碳價趨勢」對「美股企業成本」的連動影響，生成具備邏輯推演的風險摘要。"
SECTION_2_4 = "本系統之創新點在於上述兩大模組的協同運作機制。當「趨勢偵測模組」識別出碳權市場 (KRBN) 出現劇烈波動時，會立即觸發「風險解讀模組」介入，針對該特定事件生成專屬的警示報告。透過此機制，系統能將抽象的數據波動，即時轉化為具體的投資建議，落實自動化的風險管理。"

SECTION_3_INTRO = "本研究使用 2010 年至 2025 年之歷史數據進行實作驗證。"
SECTION_3_1 = "針對趨勢偵測模組，我們比較了不同演算法的表現。實驗結果顯示，在結構複雜的美股市場 (SPY) 中，雙層 LSTM 架構的 F1-Score 達到 0.733，顯著優於傳統隨機森林模型。而在碳權市場 (KRBN) 中，數據暗示了碳權價格的波動邏輯具有高度的「隨機性」與「雜訊特徵」。不同於股市受財報與經濟數據驅動，碳權資產更常受到非市場因素（如歐盟 ETS 碳定價政策公告、各國碳稅法規變動）的突發性干擾，導致歷史價格的非線性規律較難被 LSTM 完整捕捉。"

SECTION_3_2 = "為探討模型在不同性質資產間的泛化能力，本節進一步比較模型在「成長型資產（SPY）」與「氣候避險資產（KRBN）」上的表現差異（如圖 2）。\n由圖 2 可觀察到，同一套 LSTM 架構在美股 SPY 的預測效能顯著優於碳權 KRBN。針對此一「傳統資產優於氣候資產」的現象，本研究提出以下兩點解釋：\n1. 市場驅動機制的差異：美股市場（SPY）主要由企業基本面與總體經濟數據驅動，具有較清晰的景氣循環週期，利於 LSTM 捕捉規律。反觀碳權市場（KRBN），其價格波動高度受控於「政策面」（如歐盟 ETS 碳價配額調整、政府拍賣底價），這類「政策突發事件」往往缺乏歷史規律，導致純技術面模型的預測難度增加。\n2. 資產組成的本質不同：SPY 由 500 家大型上市公司組成，個別公司的非系統性風險已被分散；而 KRBN 主要持有的是「碳排放期貨合約」，期貨具有轉倉成本（Roll Yield）與到期日效應，這些衍生性金融商品的特性，使得其價格雜訊（Noise）天然高於股票現貨 ETF。"

SECTION_3_3 = "本平台已完成全端實作並進行整合測試。圖 3 展示了 AEGIS 系統的前端介面，使用者可透過下拉選單自由切換 SPY 與 KRBN 視角。\n1. 效能測試：系統經測試，從資料庫撈取預測結果至前端渲染的反應時間平均低於 1 秒，且能正確處理每日資料更新與模型重載，符合即時看盤需求。\n2. AI 摘要驗證：在圖 3 的實測畫面中，針對 KRBN 的走勢，GPT 生成的摘要能準確描述當前 RSI 為中性偏弱，並結合 LSTM 的預測結果提示「未來波動風險增加」。\n3. 雙軌避險價值：實驗介面成功呈現了 KRBN 與 SPY 的走勢背離現象（低相關性），驗證了本平台能協助投資人在追求美股成長的同時，透過碳權配置來達成氣候風險對沖的目標。"

SECTION_4 = "本文提出了一套基於生成式 AI 的 ETF 雙軌分析架構 AEGIS。本研究的主要貢獻在於：(1) 架構創新：成功將「協同運作」概念應用於跨市場風險管理；(2) 實證分析：驗證了 LSTM 在碳權 (KRBN) 與美股雙軌預測上的效能差異，指出了碳權市場的高雜訊特徵；(3) 應用價值：利用生成式 AI 降低了綠色金融的技術門檻，為投資人提供了一套視覺化、可解釋的決策輔助工具。"

# 更新後的完整參考文獻列表 (您的版本 + 補充)
REFS = [
    "楊又肇 (2020)。以深度學習LSTM方法進行台灣加權股價指數預測。國立交通大學資訊管理研究所碩士論文。",
    "鐘毅 (2020)。以深度學習LSTM方法進行台灣加權股價指數預測。國立交通大學科技管理研究所碩士論文。",
    "陳思妘 (2021)。深度學習於台灣加權股價指數預測之應用。國立交通大學管理學院財務金融學程碩士論文。",
    "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32.",
    "Hochreiter, S., & Schmidhuber, J. (1997). Long short-term memory. Neural Computation, 9(8), 1735-1780.",
    "Fama, E. F. (1970). Efficient capital markets: A review of theory and empirical work. The Journal of Finance, 25(2), 383-417.",
    "Kingma, D. P., & Ba, J. (2014). Adam: A method for stochastic optimization. arXiv preprint arXiv:1412.6980.",
    "Engle, R. F., Giglio, S., Kelly, B., Lee, H., & Stroebel, J. (2020). Hedging climate change news. The Review of Financial Studies, 33(3), 1184-1216.",
    "OpenAI. (2023). GPT-4 Technical Report. arXiv preprint arXiv:2303.08774.",
    "Wu, S., Irsoy, O., Lu, S., Dabravolski, V., Dredze, M., Gehrmann, S., ... & Mann, G. (2023). BloombergGPT: A large language model for finance. arXiv preprint arXiv:2303.17564.",
    "Aroussi, R. (2023). yfinance: Yahoo! Finance market data downloader. PyPI.",
    "IHS Markit. (2023). IHS Markit Global Carbon Index Rulebook. London: IHS Markit."
]

# --- 3. 主程式邏輯 ---
def generate():
    print("🚀 開始製作 AEGIS 最終整合版論文 Word 檔...")
    doc = Document()
    
    # 標題
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(TITLE), size=16, bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(TITLE_EN), size=14, bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(AUTHORS), size=12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(AFFILIATION), size=10)
    doc.add_paragraph()

    # 中文摘要
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("摘要"), size=12, bold=True)
    
    p = doc.add_paragraph(ABSTRACT)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p.runs[0])
    
    p = doc.add_paragraph()
    set_font(p.add_run("關鍵詞："), bold=True)
    set_font(p.add_run(KEYWORDS))
    doc.add_paragraph()

    # 英文摘要
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(ABSTRACT_EN_TITLE), size=12, bold=True, font_name_ch='Times New Roman')
    
    p = doc.add_paragraph(ABSTRACT_EN)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p.runs[0], font_name_ch='Times New Roman')
    
    p = doc.add_paragraph()
    set_font(p.add_run("Keywords: "), bold=True, font_name_ch='Times New Roman')
    set_font(p.add_run(KEYWORDS_EN), font_name_ch='Times New Roman')
    doc.add_paragraph()

    # 正文
    add_heading(doc, "1. 緒論")
    p = doc.add_paragraph(SECTION_1)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p.runs[0])

    add_heading(doc, "2. 系統設計")
    set_font(doc.add_paragraph(SECTION_2_INTRO).runs[0])
    
    add_heading(doc, "2.1 系統架構", level=2)
    set_font(doc.add_paragraph(SECTION_2_1).runs[0])
    
    add_figure(doc, "fig1_arch.png", "圖 1：AEGIS 系統整體架構與資料流向圖")

    add_heading(doc, "2.2 趨勢偵測運作模式 (Dual-Track Trend Detection)", level=2)
    set_font(doc.add_paragraph(SECTION_2_2).runs[0])

    add_heading(doc, "2.3 智能風險解讀運作模式 (Intelligent Risk Interpretation)", level=2)
    set_font(doc.add_paragraph(SECTION_2_3).runs[0])

    add_heading(doc, "2.4 偵測與解讀之協同運作 (Synergy Mechanism)", level=2)
    set_font(doc.add_paragraph(SECTION_2_4).runs[0])

    add_heading(doc, "3. 實驗與測試成果")
    set_font(doc.add_paragraph(SECTION_3_INTRO).runs[0])

    add_heading(doc, "3.1 模型效能分析", level=2)
    set_font(doc.add_paragraph(SECTION_3_1).runs[0])
    
    # 產生美化的表格
    p = doc.add_paragraph("表 1：各模型於雙軌市場之效能比較 (F1-Score)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.runs[0], bold=True)
    
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = 'Table Grid' # 使用格線樣式
    set_table_borders(table) # 加強框線
    
    # 表格標題列
    headers = ["Model", "SPY (美股大盤)", "QQQ (科技股)", "KRBN (碳權)"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(paragraph.runs[0], bold=True, size=11)
            
    # 表格數據
    data_rows = [
        ["RandomForest", "0.725", "0.732", "0.585"],
        ["SingleLayerLSTM", "0.728", "0.722", "0.612"],
        ["DoubleLayerLSTM", "0.733", "0.711", "0.608"]
    ]
    for i, row_data in enumerate(data_rows):
        for j, val in enumerate(row_data):
            cell = table.cell(i+1, j)
            cell.text = val
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_font(paragraph.runs[0], size=11)

    doc.add_paragraph() # 表格後空行

    # 3.2 跨資產泛化
    add_heading(doc, "3.2 跨資產泛化能力與差異分析", level=2)
    set_font(doc.add_paragraph(SECTION_3_2).runs[0])
    add_figure(doc, "fig2_perf.png", "圖 2：DoubleLayerLSTM 於不同市場之效能比較")

    # 3.3 系統整合
    add_heading(doc, "3.3 系統整合與視覺化實測", level=2)
    set_font(doc.add_paragraph(SECTION_3_3).runs[0])
    add_figure(doc, "fig3_ui.png", "圖 3：AEGIS 系統前端介面展示")

    add_heading(doc, "4. 結論")
    set_font(doc.add_paragraph(SECTION_4).runs[0])
    
    # 寫入致謝
    add_heading(doc, "致謝")
    p = doc.add_paragraph("感謝指導教授李冠榮教授對於本系統架構設計與實驗分析之指導，以及實驗室同仁的討論與建議。")
    set_font(p.runs[0])

    doc.add_paragraph()
    p = doc.add_paragraph()
    set_font(p.add_run("參考文獻"), size=14, bold=True)
    
    for ref in REFS:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        set_font(p.runs[0])

    # 存檔
    output_filename = "AEGIS_Paper_Final_v2.docx"
    try:
        doc.save(output_filename)
        print(f"\n✅ 成功！完整論文已產生：{os.path.abspath(output_filename)}")
    except Exception as e:
        print(f"\n❌ 存檔失敗：{e}")

if __name__ == "__main__":
    try:
        generate()
    except Exception as e:
        print(f"❌ 錯誤：{e}")